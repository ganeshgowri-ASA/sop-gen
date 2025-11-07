"""
Document Export Module
Handles export to multiple formats: DOCX, PDF, HTML, Excel
"""

import io
import os
from typing import Optional
from datetime import datetime
from .models import Document, Section


class DocumentExporter:
    """Handles document export to various formats"""

    def __init__(self):
        self.supported_formats = ['docx', 'pdf', 'html', 'excel', 'markdown']

    def export_document(self, document: Document, format: str = 'docx') -> bytes:
        """
        Export document to specified format

        Args:
            document: Document to export
            format: Output format (docx, pdf, html, excel, markdown)

        Returns:
            Bytes of exported document
        """
        format = format.lower()

        if format == 'docx':
            return self.to_docx(document)
        elif format == 'pdf':
            return self.to_pdf(document)
        elif format == 'html':
            return self.to_html(document).encode('utf-8')
        elif format == 'excel':
            return self.to_excel(document)
        elif format == 'markdown':
            return self.to_markdown(document).encode('utf-8')
        else:
            raise ValueError(f"Unsupported format: {format}")

    def to_markdown(self, doc: Document) -> str:
        """Convert document to Markdown format"""
        md_lines = []

        # Title
        md_lines.append(f"# {doc.title}\n")

        # Document metadata
        if doc.doc_number:
            md_lines.append(f"**Document Number:** {doc.doc_number}\n")

        if doc.metadata:
            if doc.metadata.get('company'):
                md_lines.append(f"**Company:** {doc.metadata['company']}\n")
            if doc.metadata.get('revision'):
                md_lines.append(f"**Revision:** {doc.metadata['revision']}\n")
            if doc.metadata.get('effective_date'):
                md_lines.append(f"**Effective Date:** {doc.metadata['effective_date']}\n")

        md_lines.append("\n---\n\n")

        # Sections
        for sec in doc.sections:
            if not sec.content or not sec.content.strip():
                continue

            # Section heading
            md_lines.append(f"## {sec.title}\n")

            # Section content based on type
            if sec.content_type == "text":
                md_lines.append(sec.content)
            elif sec.content_type == "image":
                md_lines.append(f"![{sec.title}]({sec.content})")
            elif sec.content_type == "table":
                md_lines.append(sec.content)
            elif sec.content_type == "flowchart":
                md_lines.append(f"```mermaid\n{sec.content}\n```")
            elif sec.content_type == "latex":
                md_lines.append(f"$$\n{sec.content}\n$$")

            md_lines.append("\n\n")

        # Footer
        md_lines.append("---\n")
        md_lines.append(f"\n*Document generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
        if doc.created_by:
            md_lines.append(f"*Created by: {doc.created_by}*\n")

        return "\n".join(md_lines)

    def to_html(self, doc: Document) -> str:
        """Convert document to HTML format"""
        # Convert markdown to HTML first
        md_text = self.to_markdown(doc)

        # Basic markdown to HTML conversion (you can use markdown2 library for better conversion)
        try:
            import markdown2
            html_body = markdown2.markdown(
                md_text,
                extras=["tables", "fenced-code-blocks", "break-on-newline"]
            )
        except ImportError:
            # Fallback if markdown2 not available
            html_body = md_text.replace('\n', '<br>\n')

        # Add logo and images sections
        logo_html = ""
        if doc.metadata.get('company_logo'):
            logo_data = doc.metadata['company_logo']
            if 'base64' in logo_data:
                logo_html = f'<div style="text-align: center; margin-bottom: 20px;"><img src="data:{logo_data["mime_type"]};base64,{logo_data["base64"]}" style="max-width: 200px;" alt="Company Logo"></div>'

        equipment_html = ""
        if doc.metadata.get('equipment_photos'):
            equipment_html = '<h2>Equipment Images</h2><div style="display: flex; flex-wrap: wrap; gap: 20px;">'
            for photo in doc.metadata['equipment_photos']:
                if 'base64' in photo:
                    equipment_html += f'<div><p><strong>{photo["name"]}</strong></p><img src="data:{photo["mime_type"]};base64,{photo["base64"]}" style="max-width: 400px;" alt="{photo["name"]}"></div>'
            equipment_html += '</div>'

        flowchart_html = ""
        if doc.metadata.get('flowcharts'):
            flowchart_html = '<h2>Process Flowcharts</h2><div style="display: flex; flex-wrap: wrap; gap: 20px;">'
            for flowchart in doc.metadata['flowcharts']:
                if flowchart['format'] != 'PDF' and 'base64' in flowchart:
                    flowchart_html += f'<div><p><strong>{flowchart["name"]}</strong></p><img src="data:{flowchart["mime_type"]};base64,{flowchart["base64"]}" style="max-width: 600px;" alt="{flowchart["name"]}"></div>'
                elif flowchart['format'] == 'PDF':
                    flowchart_html += f'<p>[PDF Flowchart: {flowchart["name"]}]</p>'
            flowchart_html += '</div>'

        # Wrap in HTML template
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{doc.title}</title>
    <style>
        body {{
            font-family: Arial, Helvetica, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 40px auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
            border-bottom: 2px solid #bdc3c7;
            padding-bottom: 5px;
        }}
        h3 {{
            color: #555;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        table, th, td {{
            border: 1px solid #ddd;
        }}
        th {{
            background-color: #3498db;
            color: white;
            padding: 12px;
            text-align: left;
        }}
        td {{
            padding: 10px;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 5px;
            border-radius: 3px;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        .metadata {{
            background-color: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #bdc3c7;
            font-size: 0.9em;
            color: #7f8c8d;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
    </style>
</head>
<body>
    {logo_html}
    {html_body}
    {equipment_html}
    {flowchart_html}
</body>
</html>"""
        return html

    def to_pdf(self, doc: Document) -> bytes:
        """Convert document to PDF format"""
        try:
            # Try using pdfkit (requires wkhtmltopdf)
            import pdfkit
            html_content = self.to_html(doc)
            pdf_bytes = pdfkit.from_string(html_content, False)
            return pdf_bytes

        except (ImportError, OSError):
            # Fallback: try using weasyprint
            try:
                from weasyprint import HTML
                html_content = self.to_html(doc)
                pdf_bytes = HTML(string=html_content).write_pdf()
                return pdf_bytes

            except ImportError:
                # If no PDF library available, return HTML as bytes with a note
                html_content = self.to_html(doc)
                note = "<!-- PDF generation requires pdfkit or weasyprint library -->\n"
                return (note + html_content).encode('utf-8')

    def to_docx(self, doc: Document) -> bytes:
        """Generate a Word .docx file from the document"""
        try:
            from docx import Document as WordDocument
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            word_doc = WordDocument()

            # Add company logo if available
            if doc.metadata.get('company_logo'):
                try:
                    import base64
                    logo_data = doc.metadata['company_logo']
                    if 'base64' in logo_data:
                        # Decode base64 image
                        img_bytes = base64.b64decode(logo_data['base64'])
                        img_stream = io.BytesIO(img_bytes)
                        # Add logo at the top
                        logo_para = word_doc.add_paragraph()
                        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        logo_run = logo_para.add_run()
                        logo_run.add_picture(img_stream, width=Inches(2))
                        word_doc.add_paragraph()  # Add spacing
                except Exception as e:
                    # If logo fails, continue without it
                    pass

            # Add title
            title = word_doc.add_heading(doc.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add document metadata
            if doc.doc_number or doc.metadata:
                info_table = word_doc.add_table(rows=0, cols=2)
                info_table.style = 'Light Grid Accent 1'

                if doc.doc_number:
                    row = info_table.add_row()
                    row.cells[0].text = 'Document Number'
                    row.cells[1].text = doc.doc_number

                if doc.metadata:
                    for key, value in doc.metadata.items():
                        if value and key not in ['standards', 'description']:
                            row = info_table.add_row()
                            row.cells[0].text = key.replace('_', ' ').title()
                            row.cells[1].text = str(value)

                word_doc.add_paragraph()  # Spacing

            # Add sections
            for sec in doc.sections:
                if not sec.content or not sec.content.strip():
                    continue

                # Add section heading
                word_doc.add_heading(sec.title, level=1)

                # Add content based on type
                if sec.content_type == "text":
                    word_doc.add_paragraph(sec.content)

                elif sec.content_type == "image":
                    try:
                        if os.path.exists(sec.content):
                            word_doc.add_picture(sec.content, width=Inches(6))
                        else:
                            word_doc.add_paragraph(f"[Image: {sec.content}]")
                    except Exception as e:
                        word_doc.add_paragraph(f"[Image: {sec.content} - Error loading]")

                elif sec.content_type == "table":
                    # Try to parse table content
                    self._add_table_to_word(word_doc, sec.content)

                elif sec.content_type == "flowchart":
                    word_doc.add_paragraph(f"[Flowchart: {sec.title}]")
                    word_doc.add_paragraph(sec.content, style='Code')

                elif sec.content_type == "latex":
                    word_doc.add_paragraph(f"Equation:")
                    word_doc.add_paragraph(sec.content, style='Code')

                # Add spacing
                word_doc.add_paragraph()

            # Add equipment photos section if available
            if doc.metadata.get('equipment_photos'):
                try:
                    import base64
                    word_doc.add_heading('Equipment Images', level=1)

                    equipment_photos = doc.metadata['equipment_photos']
                    for photo in equipment_photos:
                        if 'base64' in photo:
                            # Decode base64 image
                            img_bytes = base64.b64decode(photo['base64'])
                            img_stream = io.BytesIO(img_bytes)

                            # Add photo with caption
                            word_doc.add_paragraph(f"Figure: {photo['name']}")
                            photo_para = word_doc.add_paragraph()
                            photo_run = photo_para.add_run()
                            photo_run.add_picture(img_stream, width=Inches(4))
                            word_doc.add_paragraph()  # Add spacing
                except Exception as e:
                    # If photos fail, continue without them
                    pass

            # Add flowcharts section if available
            if doc.metadata.get('flowcharts'):
                try:
                    import base64
                    word_doc.add_heading('Process Flowcharts', level=1)

                    flowcharts = doc.metadata['flowcharts']
                    for flowchart in flowcharts:
                        if flowchart['format'] != 'PDF' and 'base64' in flowchart:
                            # Decode base64 image
                            img_bytes = base64.b64decode(flowchart['base64'])
                            img_stream = io.BytesIO(img_bytes)

                            # Add flowchart with caption
                            word_doc.add_paragraph(f"Flowchart: {flowchart['name']}")
                            flowchart_para = word_doc.add_paragraph()
                            flowchart_run = flowchart_para.add_run()
                            flowchart_run.add_picture(img_stream, width=Inches(5))
                            word_doc.add_paragraph()  # Add spacing
                        elif flowchart['format'] == 'PDF':
                            # For PDFs, just add a reference
                            word_doc.add_paragraph(f"[PDF Flowchart: {flowchart['name']}]")
                except Exception as e:
                    # If flowcharts fail, continue without them
                    pass

            # Add footer
            footer_para = word_doc.add_paragraph()
            footer_para.add_run(
                f"\nGenerated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            ).italic = True

            # Save to bytes
            buffer = io.BytesIO()
            word_doc.save(buffer)
            buffer.seek(0)
            return buffer.read()

        except ImportError:
            # If python-docx not available, return markdown as bytes
            return self.to_markdown(doc).encode('utf-8')

    def _add_table_to_word(self, word_doc, table_content: str):
        """Helper to add table to Word document"""
        try:
            from docx import Document as WordDocument

            lines = [line.strip() for line in table_content.split('\n') if line.strip()]

            if not lines:
                word_doc.add_paragraph(table_content)
                return

            # Check if it's markdown table
            if lines[0].startswith('|'):
                # Parse markdown table
                headers = [h.strip() for h in lines[0].strip('|').split('|')]
                data_rows = []

                for line in lines[2:]:  # Skip header and separator
                    if line.startswith('|'):
                        cells = [c.strip() for c in line.strip('|').split('|')]
                        data_rows.append(cells)

            else:
                # Assume CSV
                headers = [h.strip() for h in lines[0].split(',')]
                data_rows = [[c.strip() for c in line.split(',')] for line in lines[1:]]

            # Create table
            if headers and data_rows:
                table = word_doc.add_table(rows=1 + len(data_rows), cols=len(headers))
                table.style = 'Light Grid Accent 1'

                # Headers
                for j, header in enumerate(headers):
                    table.cell(0, j).text = header

                # Data
                for i, row in enumerate(data_rows, start=1):
                    for j, cell_text in enumerate(row):
                        if j < len(headers):
                            table.cell(i, j).text = cell_text
            else:
                word_doc.add_paragraph(table_content)

        except Exception as e:
            word_doc.add_paragraph(table_content)

    def to_excel(self, doc: Document) -> bytes:
        """Export document tables to Excel format"""
        try:
            import xlsxwriter

            output = io.BytesIO()
            workbook = xlsxwriter.Workbook(output, {'in_memory': True})

            # Create overview sheet
            overview = workbook.add_worksheet('Overview')
            bold = workbook.add_format({'bold': True})

            overview.write('A1', 'SOP Title', bold)
            overview.write('B1', doc.title)
            overview.write('A2', 'Document Number', bold)
            overview.write('B2', doc.doc_number)
            overview.write('A3', 'Created', bold)
            overview.write('B3', doc.created_at.strftime('%Y-%m-%d'))

            # Add sheets for each table section
            sheet_count = 0
            for sec in doc.sections:
                if sec.content_type == "table" and sec.content:
                    sheet_count += 1
                    sheet_name = f"{sec.title[:25]}_{sheet_count}" if len(sec.title) > 25 else sec.title[:31]
                    worksheet = workbook.add_worksheet(sheet_name)

                    # Parse table content
                    lines = [line.strip() for line in sec.content.split('\n') if line.strip()]

                    if lines and lines[0].startswith('|'):
                        # Markdown table
                        headers = [h.strip() for h in lines[0].strip('|').split('|')]
                        data_rows = []
                        for line in lines[2:]:
                            if line.startswith('|'):
                                cells = [c.strip() for c in line.strip('|').split('|')]
                                data_rows.append(cells)
                    else:
                        # CSV
                        headers = [h.strip() for h in lines[0].split(',')]
                        data_rows = [[c.strip() for c in line.split(',')] for line in lines[1:]]

                    # Write headers
                    for col, header in enumerate(headers):
                        worksheet.write(0, col, header, bold)

                    # Write data
                    for row_idx, row in enumerate(data_rows, start=1):
                        for col_idx, cell_value in enumerate(row):
                            worksheet.write(row_idx, col_idx, cell_value)

            workbook.close()
            output.seek(0)
            return output.read()

        except ImportError:
            # Fallback: return CSV
            csv_content = self._to_csv(doc)
            return csv_content.encode('utf-8')

    def _to_csv(self, doc: Document) -> str:
        """Convert document tables to CSV"""
        csv_lines = [f"# {doc.title}", f"# Document Number: {doc.doc_number}\n"]

        for sec in doc.sections:
            if sec.content_type == "table" and sec.content:
                csv_lines.append(f"\n# {sec.title}")
                csv_lines.append(sec.content)

        return '\n'.join(csv_lines)
